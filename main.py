from flask import Flask
from flask import Flask, render_template, Response, redirect, request, session, abort, url_for
from camera import VideoCamera
from camera2 import VideoCamera2
import os
import base64
from PIL import Image
from datetime import datetime
from datetime import date
import datetime
import random
from random import seed
from random import randint
from werkzeug.utils import secure_filename
from flask import send_file
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
import threading
import time
import shutil
import hashlib
import json
import cv2

import librosa
from pydub import AudioSegment
import wave

import pytesseract
from skimage.metrics import structural_similarity

#pip install python-docx
import PyPDF2
from docx import Document
import fitz

import imagehash
import PIL.Image
from PIL import Image
import urllib.request
import urllib.parse
from urllib.request import urlopen
import webbrowser

import mysql.connector

mydb = mysql.connector.connect(
  host="localhost",
  user="root",
  passwd="",
  charset="utf8",
  database="coc_new"
)


app = Flask(__name__)
##session key
app.secret_key = 'abcdef'
UPLOAD_FOLDER = 'static/upload'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
#####

@app.route('/',methods=['POST','GET'])
def index():
    cnt=0
    act=""
    msg=""
    if request.method == 'POST':
        username1 = request.form['uname']
        password1 = request.form['pass']
        mycursor = mydb.cursor()
        mycursor.execute("SELECT count(*) FROM coc_register where uname=%s && pass=%s",(username1,password1))
        myresult = mycursor.fetchone()[0]
        if myresult>0:
            session['username'] = username1
            result=" Your Logged in sucessfully**"
            return redirect(url_for('home')) 
        else:
            msg="Invalid Username or Password!"
        

    '''file = open('static/upload/E3.hash', 'rb')
    byte = file.read()
    file.close()
      
    decodeit = open('static/upload/E3.jpg', 'wb')
    decodeit.write(base64.b64decode((byte)))
    decodeit.close()'''
   

    return render_template('web/index.html',msg=msg,act=act)

@app.route('/login',methods=['POST','GET'])
def login():
    cnt=0
    act=""
    msg=""
    
    mycursor = mydb.cursor()
    if request.method == 'POST':
        username1 = request.form['uname']
        password1 = request.form['pass']
        mycursor.execute("SELECT count(*) FROM coc_login where username=%s && password=%s && utype='admin'",(username1,password1))
        myresult = mycursor.fetchone()[0]
        if myresult>0:
            session['username'] = username1
            #msg=" You are Logged in sucessfully**"
            return redirect(url_for('admin')) 
        else:
            msg="You are logged in fail!!!"
    
    return render_template('web/login.html',msg=msg,act=act)


@app.route('/login_auth',methods=['POST','GET'])
def login_auth():
    cnt=0
    act=""
    msg=""
  
    
    if request.method == 'POST':
        username1 = request.form['uname']
        password1 = request.form['pass']
        mycursor = mydb.cursor()
        mycursor.execute("SELECT count(*) FROM coc_login where username=%s && password=%s && utype='court'",(username1,password1))
        myresult = mycursor.fetchone()[0]
        if myresult>0:
            session['username'] = username1
            result=" Your Logged in sucessfully**"
            return redirect(url_for('ta_home')) 
        else:
            msg="Invalid Username or Password!"
            
    return render_template('web/login_auth.html',msg=msg,act=act)


@app.route('/admin', methods=['GET', 'POST'])
def admin():

    msg=""
    act=""
    email=""
    mess=""
    mycursor = mydb.cursor()

    mycursor.execute("SELECT count(*) FROM coc_case")
    data1 = mycursor.fetchone()[0]
    
    mycursor.execute("SELECT count(*) FROM coc_register")
    data2 = mycursor.fetchone()[0]

    mycursor.execute("SELECT count(*) FROM coc_request where status=0")
    data3 = mycursor.fetchone()[0]

    mycursor.execute("SELECT count(*) FROM coc_attack where status=0")
    data4 = mycursor.fetchone()[0]

    return render_template('admin.html', msg=msg,data1=data1,data2=data2,data3=data3,data4=data4)


#Blockchain
class Blockchain:
    def __init__(self):
        self.current_transactions = []
        self.chain = []
        self.nodes = set()

        # Create the genesis block
        self.new_block(previous_hash='1', proof=100)

    def register_node(self, address):
        """
        Add a new node to the list of nodes

        :param address: Address of node. Eg. 'http://192.168.0.5:5000'
        """

        parsed_url = urlparse(address)
        if parsed_url.netloc:
            self.nodes.add(parsed_url.netloc)
        elif parsed_url.path:
            # Accepts an URL without scheme like '192.168.0.5:5000'.
            self.nodes.add(parsed_url.path)
        else:
            raise ValueError('Invalid URL')


    def valid_chain(self, chain):
        """
        Determine if a given blockchain is valid

        :param chain: A blockchain
        :return: True if valid, False if not
        """

        last_block = chain[0]
        current_index = 1

        while current_index < len(chain):
            block = chain[current_index]
            print(f'{last_block}')
            print(f'{block}')
            print("\n-----------\n")
            # Check that the hash of the block is correct
            last_block_hash = self.hash(last_block)
            if block['previous_hash'] != last_block_hash:
                return False

            # Check that the Proof of Work is correct
            if not self.valid_proof(last_block['proof'], block['proof'], last_block_hash):
                return False

            last_block = block
            current_index += 1

        return True

    def resolve_conflicts(self):
        """
        This is our consensus algorithm, it resolves conflicts
        by replacing our chain with the longest one in the network.

        :return: True if our chain was replaced, False if not
        """

        neighbours = self.nodes
        new_chain = None

        # We're only looking for chains longer than ours
        max_length = len(self.chain)

        # Grab and verify the chains from all the nodes in our network
        for node in neighbours:
            response = requests.get(f'http://{node}/chain')

            if response.status_code == 200:
                length = response.json()['length']
                chain = response.json()['chain']

                # Check if the length is longer and the chain is valid
                if length > max_length and self.valid_chain(chain):
                    max_length = length
                    new_chain = chain

        # Replace our chain if we discovered a new, valid chain longer than ours
        if new_chain:
            self.chain = new_chain
            return True

        return False

    def new_block(self, proof, previous_hash):
        """
        Create a new Block in the Blockchain

        :param proof: The proof given by the Proof of Work algorithm
        :param previous_hash: Hash of previous Block
        :return: New Block
        """

        block = {
            'index': len(self.chain) + 1,
            'timestamp': time(),
            'transactions': self.current_transactions,
            'proof': proof,
            'previous_hash': previous_hash or self.hash(self.chain[-1]),
        }

        # Reset the current list of transactions
        self.current_transactions = []

        self.chain.append(block)
        return block

    def new_transaction(self, sender, recipient, amount):
        """
        Creates a new transaction to go into the next mined Block

        :param sender: Address of the Sender
        :param recipient: Address of the Recipient
        :param amount: Amount
        :return: The index of the Block that will hold this transaction
        """
        self.current_transactions.append({
            'sender': sender,
            'recipient': recipient,
            'amount': amount,
        })

        return self.last_block['index'] + 1

    @property
    def last_block(self):
        return self.chain[-1]

    @staticmethod
    def hash(block):
        """
        Creates a SHA-256 hash of a Block

        :param block: Block
        """

        # We must make sure that the Dictionary is Ordered, or we'll have inconsistent hashes
        block_string = json.dumps(block, sort_keys=True).encode()
        return hashlib.sha256(block_string).hexdigest()

    def proof_of_work(self, last_block):
        """
        Simple Proof of Work Algorithm:

         - Find a number p' such that hash(pp') contains leading 4 zeroes
         - Where p is the previous proof, and p' is the new proof
         
        :param last_block: <dict> last Block
        :return: <int>
        """

        last_proof = last_block['proof']
        last_hash = self.hash(last_block)

        proof = 0
        while self.valid_proof(last_proof, proof, last_hash) is False:
            proof += 1

        return proof

    @staticmethod
    def valid_proof(last_proof, proof, last_hash):
        """
        Validates the Proof

        :param last_proof: <int> Previous Proof
        :param proof: <int> Current Proof
        :param last_hash: <str> The hash of the Previous Block
        :return: <bool> True if correct, False if not.

        """

        guess = f'{last_proof}{proof}{last_hash}'.encode()
        guess_hash = hashlib.sha256(guess).hexdigest()
        return guess_hash[:4] == "0000"

def mine():
    # We run the proof of work algorithm to get the next proof...
    last_block = blockchain.last_block
    proof = blockchain.proof_of_work(last_block)

    # We must receive a reward for finding the proof.
    # The sender is "0" to signify that this node has mined a new coin.
    blockchain.new_transaction(
        sender="0",
        recipient=node_identifier,
        amount=1,
    )

    # Forge the new Block by adding it to the chain
    previous_hash = blockchain.hash(last_block)
    block = blockchain.new_block(proof, previous_hash)

    response = {
        'message': "New Block Forged",
        'index': block['index'],
        'transactions': block['transactions'],
        'proof': block['proof'],
        'previous_hash': block['previous_hash'],
    }
    return jsonify(response), 200


def new_transaction():
    values = request.get_json()

    # Check that the required fields are in the POST'ed data
    required = ['sender', 'recipient', 'amount']
    if not all(k in values for k in required):
        return 'Missing values', 400

    # Create a new Transaction
    index = blockchain.new_transaction(values['sender'], values['recipient'], values['amount'])

    response = {'message': f'Transaction will be added to Block {index}'}
    return jsonify(response), 201

def full_chain():
    response = {
        'chain': blockchain.chain,
        'length': len(blockchain.chain),
    }
    return jsonify(response), 200



def register_nodes():
    values = request.get_json()

    nodes = values.get('nodes')
    if nodes is None:
        return "Error: Please supply a valid list of nodes", 400

    for node in nodes:
        blockchain.register_node(node)

    response = {
        'message': 'New nodes have been added',
        'total_nodes': list(blockchain.nodes),
    }
    return jsonify(response), 201

def consensus():
    replaced = blockchain.resolve_conflicts()

    if replaced:
        response = {
            'message': 'Our chain was replaced',
            'new_chain': blockchain.chain
        }
    else:
        response = {
            'message': 'Our chain is authoritative',
            'chain': blockchain.chain
        }

    return jsonify(response), 200

def cocchain(uid,uname,bcdata,utype):
    ############

    now = datetime.datetime.now()
    yr=now.strftime("%Y")
    mon=now.strftime("%m")
    rdate=now.strftime("%d-%m-%Y")
    rtime=now.strftime("%H:%M:%S")
    
    ff=open("static/key.txt","r")
    k=ff.read()
    ff.close()
    
    #bcdata="CID:"+uname+",Time:"+val1+",Unit:"+val2
    dtime=rdate+","+rtime

    ff1=open("static/assets/js/d1.txt","r")
    bc1=ff1.read()
    ff1.close()
    
    px=""
    if k=="1":
        px=""
        result = hashlib.md5(bcdata.encode())
        key=result.hexdigest()
        print(key)
        v=k+"##"+key+"##"+bcdata+"##"+dtime

        ff1=open("static/assets/js/d1.txt","w")
        ff1.write(v)
        ff1.close()
        
        dictionary = {
            "ID": "1",
            "Pre-hash": "00000000000000000000000000000000",
            "Hash": key,
            "utype": utype,
            "Date/Time": dtime
        }

        k1=int(k)
        k2=k1+1
        k3=str(k2)
        ff1=open("static/key.txt","w")
        ff1.write(k3)
        ff1.close()

        ff1=open("static/prehash.txt","w")
        ff1.write(key)
        ff1.close()
        
    else:
        px=","
        pre_k=""
        k1=int(k)
        k2=k1-1
        k4=str(k2)

        ff1=open("static/prehash.txt","r")
        pre_hash=ff1.read()
        ff1.close()
        
        g1=bc1.split("#|")
        for g2 in g1:
            g3=g2.split("##")
            if k4==g3[0]:
                pre_k=g3[1]
                break

        
        result = hashlib.md5(bcdata.encode())
        key=result.hexdigest()
        

        v="#|"+k+"##"+key+"##"+bcdata+"##"+dtime

        k3=str(k2)
        ff1=open("static/key.txt","w")
        ff1.write(k3)
        ff1.close()

        ff1=open("static/assets/js/d1.txt","a")
        ff1.write(v)
        ff1.close()

        
        
        dictionary = {
            "ID": k,
            "Pre-hash": pre_hash,
            "Hash": key,
            "utype:": utype,
            "Date/Time": dtime
        }
        k21=int(k)+1
        k3=str(k21)
        ff1=open("static/key.txt","w")
        ff1.write(k3)
        ff1.close()

        ff1=open("static/prehash.txt","w")
        ff1.write(key)
        ff1.close()

    m=""
    if k=="1":
        m="w"
    else:
        m="a"
    # Serializing json
    
    json_object = json.dumps(dictionary, indent=4)
     
    # Writing to sample.json
    with open("static/cocchain.json", m) as outfile:
        outfile.write(json_object)
    ##########
        


@app.route('/view_req', methods=['GET', 'POST'])
def view_req():

    msg=""
    act=request.args.get("act")
    st=""
    rid=request.args.get("rid")
    email=""
    mess=""
    data=[]

    bdata=""
   
    
    mycursor = mydb.cursor()

    mycursor.execute("SELECT * FROM coc_register order by id")
    data2 = mycursor.fetchall()
        
    mycursor.execute("SELECT count(*) FROM coc_request")
    cnt = mycursor.fetchone()[0]
    if cnt>0:
        st="1"
        mycursor.execute("SELECT * FROM coc_request order by id desc")
        data = mycursor.fetchall()

        mycursor.execute("update coc_request set status=1")
        mydb.commit()

    
    if request.method == 'POST':
        if act=="send":
            user = request.form['user']
            message = request.form['message']
            mycursor.execute("SELECT max(id)+1 FROM coc_request")
            maxid = mycursor.fetchone()[0]
            if maxid is None:
                maxid=1
            sql = "INSERT INTO coc_request(id,uname,message,reply,status,cname) VALUES (%s,%s,%s,%s,%s,%s)"
            val = (maxid,'admin',message,'','0',user)
            mycursor.execute(sql, val)
            mydb.commit()
            ###
            mycursor.execute('SELECT * FROM coc_request WHERE id=%s', (maxid,))
            dd = mycursor.fetchone()
            dtime=str(dd[5])
            bcdata="Request ID:"+str(maxid)+", User ID:"+user+", Status:Request by admin"
            cocchain(str(maxid),user,bcdata,'Request')
            ###
            msg="send"
        else:
            
            reply = request.form['reply']
            mycursor.execute("update coc_request set reply=%s where id=%s",(reply,rid))
            mydb.commit()
            ###
            mycursor.execute('SELECT * FROM coc_request WHERE id=%s', (rid,))
            dd = mycursor.fetchone()
            dtime=str(dd[5])
            bcdata="Request ID:"+rid+", User ID:"+dd[1]+", Status:Reply by admin"
            cocchain(str(rid),dd[1],bcdata,'REply')
            ###
            msg="reply"

        
    

    return render_template('view_req.html', msg=msg,act=act,data=data,data2=data2,st=st)

@app.route('/view_noti', methods=['GET', 'POST'])
def view_noti():

    msg=""
    act=request.args.get("act")
    st=""
    email=""
    mess=""
    data=[]

    bdata=""
    f1=open("bc.txt","r")
    bc=f1.read()
    f1.close()
    
    mycursor = mydb.cursor()

    mycursor.execute("SELECT count(*) FROM coc_attack")
    cnt = mycursor.fetchone()[0]
    if cnt>0:
        st="1"
        mycursor.execute("SELECT * FROM coc_attack order by id desc")
        data = mycursor.fetchall()

        mycursor.execute("update coc_attack set status=1")
        mydb.commit()

    return render_template('view_noti.html', msg=msg,act=act,data=data,st=st)

@app.route('/contact', methods=['GET', 'POST'])
def contact():
    msg=""
    act=""
    data=[]
    email=""
    mess=""
    mycursor = mydb.cursor()

    mycursor.execute("SELECT * FROM coc_login where username='admin'")
    data = mycursor.fetchone()

    if request.method=='POST':
        email=request.form['email']
        mycursor.execute("update coc_login set email=%s where username='admin'",(email,))
        mydb.commit()
        msg="success"
    return render_template('contact.html', msg=msg,act=act,data=data)

@app.route('/add_auth', methods=['GET', 'POST'])
def add_auth():
    msg=""
    act=""
    data=[]
    email=""
    mess=""
    mycursor = mydb.cursor()
    bdata=""

    #f1=open("bc.txt","r")
    #bc=f1.read()
    #f1.close()
            
    if request.method=='POST':
        name=request.form['name']
        designation=request.form['designation']
        mobile=request.form['mobile']
        email=request.form['email']
        aadhar=request.form['aadhar']
        location=request.form['location']
        city=request.form['city']

        

        mycursor.execute("SELECT count(*) FROM coc_register where aadhar=%s",(aadhar,))
        myresult = mycursor.fetchone()[0]

        if myresult==0:
        
            mycursor.execute("SELECT max(id)+1 FROM coc_register")
            maxid = mycursor.fetchone()[0]
            if maxid is None:
                maxid=1

            uname="AT"+str(maxid)
            p1=randint(1000,9999)
            pass1="123456"
            now = date.today() #datetime.datetime.now()
            rdate=now.strftime("%d-%m-%Y")
            
            sql = "INSERT INTO coc_register(id,name,designation,mobile,email,aadhar,location,city,uname,pass,status) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
            val = (maxid,name,designation,mobile,email,aadhar,location,city,uname,pass1,'1')
            mycursor.execute(sql, val)
            mydb.commit()

            
            #print(mycursor.rowcount, "Registered Success")
            msg="success"
            act="1"
            mess="Dear "+name+", Authorized Party - User ID: "+uname+", Password: "+pass1
            ###
            mycursor.execute('SELECT * FROM coc_register WHERE id=%s', (maxid,))
            dd = mycursor.fetchone()
            dtime=str(dd[11])
            bcdata="ID:"+str(maxid)+", User ID:"+uname+", Status:Authorized User Created, Aadhar:"+aadhar
            cocchain(str(maxid),uname,bcdata,'AP')
            ###
            
        else:
            
            msg='fail'

    mycursor.execute("SELECT * FROM coc_register")
    data = mycursor.fetchall()
    
    return render_template('add_auth.html',msg=msg,email=email,mess=mess,act=act,data=data)

@app.route('/add_case', methods=['GET', 'POST'])
def add_case():
    msg=""
    act=""
    email=""
    mess=""

    bdata=""
    #f1=open("bc.txt","r")
    #bc=f1.read()
    #f1.close()
    
    mycursor = mydb.cursor()
    

    if request.method=='POST':
        district=request.form['district']
        station=request.form['station']
        title=request.form['title']
        cdate=request.form['cdate']
        details=request.form['details']
        suspect=request.form['suspect']
        name=request.form['name']
        fname=request.form['fname']
        gender=request.form['gender']
        dob=request.form['dob']
        address=request.form['address']
        district2=request.form['district2']
        pincode=request.form['pincode']
        mobile=request.form['mobile']
        email=request.form['email']
        aadhar=request.form['aadhar']
        
        mycursor.execute("SELECT max(id)+1 FROM coc_case")
        maxid = mycursor.fetchone()[0]
        if maxid is None:
            maxid=1

        now = date.today() #datetime.datetime.now()
        rdate=now.strftime("%d-%m-%Y")

        mm=now.strftime("%m")
        yy=now.strftime("%Y")
        case_id="C"+mm+yy+str(maxid)
        
        sql = "INSERT INTO coc_case(id,case_id,district,station,title,cdate,details,suspect,name,fname,gender,dob,address,district2,pincode,mobile,email,aadhar,status) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
        val = (maxid,case_id,district,station,title,cdate,details,suspect,name,fname,gender,dob,address,district2,pincode,mobile,email,aadhar,'0')
        mycursor.execute(sql, val)
        mydb.commit()

        
        print(mycursor.rowcount, "Registered Success")
        msg="success"
        act="1"
        #mess="Dear "+name+", User ID: "+uname+", Password: "+pass1
        ###
        mycursor.execute('SELECT * FROM coc_case WHERE id=%s', (maxid,))
        dd = mycursor.fetchone()
        dtime=str(dd[19])
        bcdata="ID:"+str(maxid)+", Case ID:"+case_id+", Status:Case Registered, Complainant Name: "+name
        cocchain(str(maxid),case_id,bcdata,'Case')
        ###
     
    
    return render_template('add_case.html',msg=msg,email=email,mess=mess,act=act)

##########
def hash_file(file_path):
    # Initialize a hash object
    hash_obj = hashlib.sha256()

    try:
        with open(file_path, 'rb') as file:
            # Read the file in chunks to handle large files efficiently
            chunk = 0
            while chunk != b'':
                chunk = file.read(1024)  # Read 1 KB at a time
                hash_obj.update(chunk)  # Update hash with each chunk

            # Calculate the hexadecimal digest of the hash
            file_hash = hash_obj.hexdigest()
            return file_hash
    except FileNotFoundError:
        print("File not found")
        return None

def hash_pdf(file_path):
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfFileReader(file)
            text = ""
            for page_num in range(pdf_reader.numPages):
                text += pdf_reader.getPage(page_num).extractText()
            return hashlib.sha256(text.encode()).hexdigest()
    except FileNotFoundError:
        print("File not found")
        return None

def hash_docx(file_path):
    try:
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return hashlib.sha256(text.encode()).hexdigest()
    except FileNotFoundError:
        print("File not found")
        return None

###########
@app.route('/add_evidence', methods=['GET', 'POST'])
def add_evidence():
    msg=""
    cid=request.args.get("cid")
    act=request.args.get("act")
    email=""
    mess=""
    efile=""
    data2=[]
    st=""

    bdata=""
    
    
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM coc_case where id=%s",(cid,))
    data = mycursor.fetchone()
    case_id=data[1]

    mycursor.execute("SELECT count(*) FROM coc_evidence where case_id=%s",(case_id,))
    cnt = mycursor.fetchone()[0]
    if cnt>0:
        st="1"
        mycursor.execute("SELECT * FROM coc_evidence where case_id=%s",(case_id,))
        data2 = mycursor.fetchall()    

    if request.method=='POST':
        details=request.form['details']
        file=request.files['file']
        
        
        mycursor.execute("SELECT max(id)+1 FROM coc_evidence")
        maxid = mycursor.fetchone()[0]
        if maxid is None:
            maxid=1

        now = date.today() #datetime.datetime.now()
        rdate=now.strftime("%d-%m-%Y")

        if file:
            fname = file.filename
            filename = secure_filename(fname)
            efile="E"+str(maxid)+filename
            file.save(os.path.join("static/upload1", efile))
            shutil.copy("static/upload1/"+efile,"static/assets/ud/"+efile)

            with open("static/upload1/"+efile, "rb") as image2string:
                converted_string = base64.b64encode(image2string.read())
            print(converted_string)
            bfile1="E"+str(maxid)+".hash"
            with open('static/upload/'+bfile1, "wb") as file:
                file.write(converted_string)

        mm=now.strftime("%m")
        yy=now.strftime("%Y")

        ef=efile.split(".")
        if ef[1]=="jpg" or ef[1]=="jpeg" or ef[1]=="png":
            md5hash = hashlib.md5(Image.open("static/upload1/"+efile).tobytes())
            hashval=md5hash.hexdigest()
        else:
            hashval=sha256_hash("static/upload1/"+efile)
        
        if ef[1]=="wav":
            plot_waveform(efile,str(maxid))

            
        sql = "INSERT INTO coc_evidence(id,case_id,details,filename,upload_by,hash_value) VALUES (%s,%s,%s,%s,%s,%s)"
        val = (maxid,case_id,details,efile,'admin',hashval)
        mycursor.execute(sql, val)
        mydb.commit()

        
        print(mycursor.rowcount, "Registered Success")
        msg="success"
        act="1"
        #mess="Dear "+name+", User ID: "+uname+", Password: "+pass1
        ###
        mycursor.execute('SELECT * FROM coc_evidence WHERE id=%s', (maxid,))
        dd = mycursor.fetchone()
        dtime=str(dd[4])
        bcdata="Evidence ID:"+str(maxid)+", Case ID:"+case_id+", Status: Evidence File: "+efile+", Upload by admin"
        cocchain(str(maxid),case_id,bcdata,'Evidence')
        ###
     
    if act=="del":
        did=request.args.get("did")

        ###
        mycursor.execute("update coc_evidence set status=0 where id=%s",(did,))
        mydb.commit()
        mycursor.execute('SELECT * FROM coc_evidence WHERE id=%s', (did,))
        dd = mycursor.fetchone()
        dtime=str(dd[4])
        #bdata="Evidence ID:"+str(did)+", Case ID:"+case_id+", Status:Evidence File Deleted, File: "+dd[3]+", Date: "+dtime
        
        ###
        msg="deleted"

        
        mycursor.execute("delete from coc_evidence where id=%s",(did,))
        mydb.commit()
        #return redirect(url_for('add_evidence',cid=cid))
    
    return render_template('add_evidence.html',msg=msg,email=email,mess=mess,act=act,cid=cid,case_id=case_id,data2=data2,st=st)

@app.route('/allow', methods=['GET', 'POST'])
def allow():
    msg=""
    
    cid=request.args.get("cid")
    act=request.args.get("act")
    email=""
    mess=""
    efile=""
    data2=[]
    st=""
    q=""
    s1="0"
    s2="0"

    bdata=""
    
    
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM coc_case where id=%s",(cid,))
    data = mycursor.fetchone()
    
    case_id=data[1]

    mycursor.execute("SELECT * FROM coc_register")
    data2 = mycursor.fetchall()



    if request.method=='POST':
        user=request.form['user']
        ch=request.form.getlist('ch[]')
        print(ch)
        l=len(ch)
        if l==2:
            s1="1"
            s2="1"
            q="View and Upload"
        elif l==1:
            if ch[0]=="1":
                s1="1"
                s2="0"
                q="View"
            else:
                s2="1"
                s1="1"
                q="View and Upload"

        mycursor.execute('SELECT count(*) FROM coc_allow WHERE uname=%s && case_id=%s', (user,case_id))
        c1 = mycursor.fetchone()[0]
        if c1==0:

            mycursor.execute("SELECT max(id)+1 FROM coc_allow")
            maxid = mycursor.fetchone()[0]
            if maxid is None:
                maxid=1
                
            sql = "INSERT INTO coc_allow(id,uname,case_id,view_st,upload_st) VALUES (%s,%s,%s,%s,%s)"
            val = (maxid,user,case_id,s1,s2)
            mycursor.execute(sql, val)
            mydb.commit()

            ###
            mycursor.execute('SELECT * FROM coc_allow WHERE id=%s', (maxid,))
            dd = mycursor.fetchone()
            dtime=str(dd[5])
            bcdata="Allow ID:"+str(maxid)+", Case ID:"+case_id+", Status:Allowed for "+q+", User:"+user
            cocchain(str(maxid),case_id,bcdata,'Allow')
            ###
        else:
            mycursor.execute("update coc_allow set view_st=%s,upload_st=%s where uname=%s && case_id=%s", (s1,s2,user,case_id))
            mydb.commit()
            ###
            mycursor.execute("SELECT * FROM coc_allow WHERE uname=%s && case_id=%s", (user,case_id))
            dd = mycursor.fetchone()
            dtime=str(dd[5])
            bcdata="Allow ID:"+str(dd[0])+", Case ID:"+case_id+", Status:Allowed for "+q+", User:"+user
            cocchain(str(dd[0]),case_id,bcdata,'Allow')
            ###

            
        msg="allow"
        
        
    
    return render_template('allow.html',msg=msg,data=data,act=act,cid=cid,case_id=case_id,data2=data2,st=st)

@app.route('/access', methods=['GET', 'POST'])
def access():
    msg=""
    eid=request.args.get("eid")
    cid=request.args.get("cid")
    act=request.args.get("act")
    email=""
    mess=""
    efile=""
    data2=[]
    st=""
    q=""
    s1="0"
    s2="0"

    bdata=""
    
    
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM coc_evidence where id=%s",(eid,))
    data = mycursor.fetchone()
    
    case_id=data[1]

    mycursor.execute("SELECT * FROM coc_register")
    data2 = mycursor.fetchall()



    if request.method=='POST':
        user=request.form['user']
        ch=request.form.getlist('ch[]')
        print(ch)
        l=len(ch)
        if l==2:
            s1="1"
            s2="1"
            q="View and Download"
        elif l==1:
            if ch[0]==1:
                s1="1"
                q="View"
            else:
                s2="1"
                if q=="":
                    q="Download"
                else:
                    q="View and Download"

        mycursor.execute('SELECT count(*) FROM coc_access WHERE uname=%s && eid=%s', (user,eid))
        c1 = mycursor.fetchone()[0]
        if c1==0:

            mycursor.execute("SELECT max(id)+1 FROM coc_access")
            maxid = mycursor.fetchone()[0]
            if maxid is None:
                maxid=1
                
            sql = "INSERT INTO coc_access(id,uname,eid,case_id,view_st,download_st) VALUES (%s,%s,%s,%s,%s,%s)"
            val = (maxid,user,eid,case_id,s1,s2)
            mycursor.execute(sql, val)
            mydb.commit()

            ###
            mycursor.execute('SELECT * FROM coc_access WHERE id=%s', (maxid,))
            dd = mycursor.fetchone()
            dtime=str(dd[6])
            bcdata="Access ID:"+str(maxid)+", Case ID:"+case_id+", Status:Access for "+q+", User:"+user
            cocchain(str(maxid),case_id,bcdata,'Access')
            ###
        else:
            mycursor.execute("update coc_access set view_st=%s,download_st=%s where uname=%s && case_id=%s", (s1,s2,user,case_id))
            mydb.commit()
            ###
            mycursor.execute("SELECT * FROM coc_access WHERE uname=%s && eid=%s", (user,eid))
            dd = mycursor.fetchone()
            dtime=str(dd[6])
            bcdata="Access ID:"+str(dd[0])+", Case ID:"+case_id+", Status:Access for "+q+", User:"+user
            cocchain(str(dd[0]),case_id,bcdata,'Access')
            ###

            
        msg="access"
        
    
    return render_template('access.html',msg=msg,data=data,act=act,cid=cid,case_id=case_id,data2=data2,st=st,eid=eid)



@app.route('/view_auth', methods=['GET', 'POST'])
def view_auth():
    msg=""
    act=request.args.get("act")
    email=""
    mess=""
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM coc_register")
    data = mycursor.fetchall()

    if act=="del":
        did=request.args.get("did")
        mycursor.execute("delete from coc_register where id=%s",(did,))
        mydb.commit()
        return redirect(url_for('view_auth'))
    
    return render_template('view_auth.html',msg=msg,act=act,data=data)


@app.route('/view_case', methods=['GET', 'POST'])
def view_case():
    msg=""
    act=request.args.get("act")
    email=""
    mess=""
    bdata=""
    
    
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM coc_case order by id desc")
    data = mycursor.fetchall()

    if act=="del":
        did=request.args.get("did")
        
        ###
        mycursor.execute("update coc_case set status=0 where id=%s",(did,))
        mydb.commit()
        mycursor.execute('SELECT * FROM coc_case WHERE id=%s', (did,))
        dd = mycursor.fetchone()
        dtime=str(dd[19])
        #bdata="ID:"+str(did)+", Case ID:"+dd[1]+", Status:Case Deleted, Date: "+dtime
        ###
        mycursor.execute("delete from coc_evidence where case_id=%s",(dd[1],))
        mydb.commit()
        mycursor.execute("delete from coc_case where id=%s",(did,))
        mydb.commit()
        msg="deleted"
        #return redirect(url_for('view_case'))
    
    return render_template('view_case.html',msg=msg,act=act,data=data)


@app.route('/home', methods=['GET', 'POST'])
def home():
    msg=""
    act=""
    data=[]
    uname=""
    if 'username' in session:
        uname = session['username']

    print(uname)
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM coc_register where uname=%s",(uname, ))
    data1 = mycursor.fetchone()

    mycursor.execute("SELECT count(*) FROM coc_case c,coc_allow a where a.case_id=c.case_id && a.uname=%s order by c.id desc",(uname,))
    cnt = mycursor.fetchone()[0]
    if cnt>0:
        st="1"
        mycursor.execute("SELECT * FROM coc_case c,coc_allow a where a.case_id=c.case_id && a.uname=%s order by c.id desc",(uname,))
        data = mycursor.fetchall()

    
    return render_template('home.html',data=data,data1=data1,act=act)

@app.route('/a_view_case', methods=['GET', 'POST'])
def a_view_case():
    msg=""
    act=request.args.get("act")
    uname=""
    if 'username' in session:
        uname = session['username']

    data=[]
    email=""
    mess=""
    bdata=""
    f1=open("bc.txt","r")
    bc=f1.read()
    f1.close()

    st=""
    
    mycursor = mydb.cursor()
    mycursor.execute("SELECT count(*) FROM coc_case c,coc_allow a where a.case_id=c.case_id && a.uname=%s order by c.id desc",(uname,))
    cnt = mycursor.fetchone()[0]
    if cnt>0:
        st="1"
        mycursor.execute("SELECT * FROM coc_case c,coc_allow a where a.case_id=c.case_id && a.uname=%s order by c.id desc",(uname,))
        data = mycursor.fetchall()

    
    
    return render_template('a_view_case.html',msg=msg,act=act,data=data,bc=bc,bdata=bdata)

################
def calculate_hash(file_path):
    # Calculate the hash value of a file
    hasher = hashlib.sha256()
    with open(file_path, 'rb') as f:
        while True:
            data = f.read(65536)  # Read the file in chunks to avoid loading it entirely into memory
            if not data:
                break
            hasher.update(data)
    return hasher.hexdigest()

def hash_difference_percentage(hash1, hash2):
    # Calculate the percentage difference between two hash values
    if len(hash1) != len(hash2):
        raise ValueError("Hash values must have the same length")
    
    difference_count = sum(c1 != c2 for c1, c2 in zip(hash1, hash2))
    total_length = len(hash1)
    percentage_difference = (difference_count / total_length) * 100
    return percentage_difference


def sha256_hash(file_path):
    try:
        with open(file_path, 'rb') as file:
            # Read the entire file
            data = file.read()
            # Calculate the SHA-256 hash
            sha256_hash = hashlib.sha256(data).hexdigest()
            return sha256_hash
    except FileNotFoundError:
        print("File not found")
        return None

def tamper_video(file1,file2,vtype):

    ###
    folder = 'static/video1'
    for filename in os.listdir(folder):
        file_path = os.path.join(folder, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            print('Failed to delete %s. Reason: %s' % (file_path, e))
    ###
    folder = 'static/video2'
    for filename in os.listdir(folder):
        file_path = os.path.join(folder, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            print('Failed to delete %s. Reason: %s' % (file_path, e))
    ###
    folder = 'static/video3'
    for filename in os.listdir(folder):
        file_path = os.path.join(folder, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            print('Failed to delete %s. Reason: %s' % (file_path, e))
    ###
    video_path1=""
    #vtype=1 tamper, vtype=2 attack
    if vtype=="1":
        video_path1="static/assets/ud/"+file1
    else:
        video_path1="static/assets/ud/"+file1
        
    video_path2="static/upload1/"+file2
    cap1 = cv2.VideoCapture(video_path1)
    cap2 = cv2.VideoCapture(video_path2)

    frame_number = 1
    tampering_frames = []
    try:
        while True:
            ret1, frame1 = cap1.read()
            ret2, frame2 = cap2.read()

            fr1="a"+str(frame_number)+".jpg"
            fr2="b"+str(frame_number)+".jpg"
            print(fr1)
            cv2.imwrite("static/video1/"+fr1, frame1)
            cv2.imwrite("static/video2/"+fr2, frame2)
                
            

            frame_number += 1
    except:
        print("try")
        frame_number=5
    ###
    i=1
    while i<frame_number:
        gr1="a"+str(i)+".jpg"
        gr2="b"+str(i)+".jpg"
        before = cv2.imread('static/video1/'+gr1)
        after = cv2.imread('static/video2/'+gr2)

        # Convert images to grayscale
        before_gray = cv2.cvtColor(before, cv2.COLOR_BGR2GRAY)
        after_gray = cv2.cvtColor(after, cv2.COLOR_BGR2GRAY)

        # Compute SSIM between the two images
        (score, diff) = structural_similarity(before_gray, after_gray, full=True)
        print("Image Similarity: {:.4f}%".format(score * 100))
        per=format(score * 100)

        # The diff image contains the actual image differences between the two images
        # and is represented as a floating point data type in the range [0,1] 
        # so we must convert the array to 8-bit unsigned integers in the range
        # [0,255] before we can use it with OpenCV
        diff = (diff * 255).astype("uint8")
        diff_box = cv2.merge([diff, diff, diff])

        # Threshold the difference image, followed by finding contours to
        # obtain the regions of the two input images that differ
        thresh = cv2.threshold(diff, 0, 255, cv2.THRESH_BINARY_INV | cv2.THRESH_OTSU)[1]
        contours = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        contours = contours[0] if len(contours) == 2 else contours[1]

        mask = np.zeros(before.shape, dtype='uint8')
        filled_after = after.copy()
        j=1
        for c in contours:
            area = cv2.contourArea(c)
            if area > 40:
                x,y,w,h = cv2.boundingRect(c)
                cv2.rectangle(before, (x, y), (x + w, y + h), (36,255,12), 2)
                mm=cv2.rectangle(after, (x, y), (x + w, y + h), (36,255,12), 2)
                cv2.imwrite("static/video3/"+gr2, mm)
        i+=1



@app.route('/a_upload', methods=['GET', 'POST'])
def a_upload():
    msg=""
    cid=request.args.get("cid")
    eid=request.args.get("eid")
    act=request.args.get("act")
    email=""
    mess=""
    efile=""
    data1=[]
    data2=[]
    bdata1=""
    msg1=""
    st=""
    file1=""
    file2=""
    pfile=""
    fid=0

    uname=""
    if 'username' in session:
        uname = session['username']

    bdata=""
    
    mycursor = mydb.cursor()

    mycursor.execute("SELECT * FROM coc_register where uname=%s",(uname,))
    value = mycursor.fetchone()

    mycursor.execute("SELECT * FROM coc_login where username='admin'")
    avalue = mycursor.fetchone()
    email=avalue[4]
    
    mycursor.execute("SELECT * FROM coc_case where id=%s",(cid,))
    data = mycursor.fetchone()
    case_id=data[1]

    mycursor.execute("SELECT count(*) FROM coc_case c,coc_allow a where a.case_id=c.case_id && a.uname=%s order by c.id desc",(uname,))
    cnt1 = mycursor.fetchone()[0]
    if cnt1>0:
        st="1"
        mycursor.execute("SELECT * FROM coc_case c,coc_allow a where a.case_id=c.case_id && a.uname=%s && a.case_id=%s order by c.id desc",(uname,case_id))
        data1 = mycursor.fetchone()

        

    mycursor.execute("SELECT count(*) FROM coc_evidence where case_id=%s",(case_id,))
    cnt = mycursor.fetchone()[0]
    if cnt>0:
        st="1"
        mycursor.execute("SELECT * FROM coc_evidence where case_id=%s",(case_id,))
        dd2 = mycursor.fetchall()
        for ds2 in dd2:
            dt2=[]
            dt2.append(ds2[0])
            dt2.append(ds2[1])
            dt2.append(ds2[2])
            dt2.append(ds2[3])
            dt2.append(ds2[4])
            dt2.append(ds2[5])
            dt2.append(ds2[6])

            mycursor.execute("SELECT count(*) FROM coc_access where eid=%s && uname=%s",(ds2[0],uname))
            n3 = mycursor.fetchone()[0]
            if n3>0:
                mycursor.execute("SELECT * FROM coc_access where eid=%s && uname=%s",(ds2[0],uname))
                d3 = mycursor.fetchone()
                if d3[5]==1:
                    dt2.append("1")
                elif d3[5]==2:
                    dt2.append("2")
                else:
                    dt2.append("3")
            else:
                dt2.append("3")

            data2.append(dt2)
            

    if request.method=='POST':
        details=request.form['details']
        file=request.files['file']
        
        
        mycursor.execute("SELECT max(id)+1 FROM coc_evidence")
        maxid = mycursor.fetchone()[0]
        if maxid is None:
            maxid=1

        now = date.today() #datetime.datetime.now()
        rdate=now.strftime("%d-%m-%Y")

        if file:
            fname = file.filename
            filename = secure_filename(fname)
            efile="E"+str(maxid)+filename
            file.save(os.path.join("static/upload1", efile))
            shutil.copy("static/upload1/"+efile,"static/assets/ud/"+efile)

            with open("static/upload1/"+efile, "rb") as image2string:
                converted_string = base64.b64encode(image2string.read())
            #print(converted_string)
            bfile1="E"+str(maxid)+".hash"
            with open('static/upload/'+bfile1, "wb") as file:
                file.write(converted_string)

        mm=now.strftime("%m")
        yy=now.strftime("%Y")

        ef=efile.split(".")
       
        if ef[1]=="jpg" or ef[1]=="jpeg" or ef[1]=="png":
            md5hash = hashlib.md5(Image.open("static/upload1/"+efile).tobytes())
            hashval=md5hash.hexdigest()
        else:
            hashval=sha256_hash("static/upload1/"+efile)
            
        #md5hash = hashlib.md5(Image.open("static/upload1/"+efile).tobytes())
        #hashval=md5hash.hexdigest()
        
        sql = "INSERT INTO coc_evidence(id,case_id,details,filename,upload_by,hash_value) VALUES (%s,%s,%s,%s,%s,%s)"
        val = (maxid,case_id,details,efile,uname,hashval)
        mycursor.execute(sql, val)
        mydb.commit()

        
        print(mycursor.rowcount, "Registered Success")
        #msg="success"
        act="1"
        #mess="Dear "+name+", User ID: "+uname+", Password: "+pass1
        ###
        mycursor.execute('SELECT * FROM coc_evidence WHERE id=%s', (maxid,))
        dd = mycursor.fetchone()
        dtime=str(dd[4])
        bcdata="Evidence ID:"+str(maxid)+", Case ID:"+case_id+", Status: Evidence File: "+efile+", upload by "+uname
        cocchain(str(maxid),case_id,bcdata,'Upload')
        #################
        
        selected_fn=""
        ####Fuzzy hash similarity verification######

        file_ext=efile.split(".")
        

        mycursor.execute('SELECT * FROM coc_evidence WHERE id!=%s && case_id=%s && tamper_st=0', (maxid,case_id))
        drow = mycursor.fetchall()

        x=0
        for drow1 in drow:
            pfile=drow1[3]
            fid=drow1[0]

            fext=pfile.split(".")
            if fext[1]==file_ext[1]:               

                if file_ext[1]=="jpg" or file_ext[1]=="jpeg" or file_ext=="png":
                    d=1

                    hash1 = calculate_hash("static/upload1/"+pfile)
                    hash2 = calculate_hash("static/upload1/"+efile)

                    percentage_difference = hash_difference_percentage(hash1, hash2)
                    if percentage_difference>=95:
                        selected_fn=pfile
                        st="no"
                        x+=1
                        mycursor.execute("update coc_evidence set tamper_st=1,ftype='image',fname=%s,fid=%s where id=%s",(pfile,fid,maxid))
                        mydb.commit()
                        break
                        
                        
                    
                elif file_ext[1]=="docx":
                    d=1
                    hash1 = calculate_hash("static/upload1/"+pfile)
                    hash2 = calculate_hash("static/upload1/"+efile)

                    percentage_difference = hash_difference_percentage(hash1, hash2)
                    print("docx diff")
                    print(percentage_difference)
                    if percentage_difference>=95:
                        selected_fn=pfile
                        st="docx"
                        x+=1
                        mycursor.execute("update coc_evidence set tamper_st=1,ftype='docx',fname=%s,fid=%s where id=%s",(pfile,fid,maxid))
                        mydb.commit()
                        break
                elif file_ext[1]=="pdf":
                    d=1
                    hash1 = calculate_hash("static/upload1/"+pfile)
                    hash2 = calculate_hash("static/upload1/"+efile)
                    
                    percentage_difference = hash_difference_percentage(hash1, hash2)
                    print("pdf diff")
                    print(percentage_difference)
                    if percentage_difference>=95:
                        selected_fn=pfile
                        st="pdf"
                        x+=1
                        mycursor.execute("update coc_evidence set tamper_st=1,ftype='pdf',fname=%s,fid=%s where id=%s",(pfile,fid,maxid))
                        mydb.commit()
                        break
                elif file_ext[1]=="wav" or file_ext[1]=="mp3":
                    d=1
                    plot_waveform(efile,str(maxid))
                    res=compare_audio("static/upload1/"+pfile,"static/upload1/"+efile)
                    print("audio per")
                    print(res[2])
                    if res[2]>0 and res[2]<3:
                        
                        st="audio"
                        x+=1
                        mycursor.execute("update coc_evidence set tamper_st=1,ftype='audio',fname=%s,fid=%s where id=%s",(pfile,fid,maxid))
                        mydb.commit()
                        break
                    '''hash1 = calculate_hash("static/upload1/"+pfile)
                    hash2 = calculate_hash("static/upload1/"+efile)

                    percentage_difference = hash_difference_percentage(hash1, hash2)
                    if percentage_difference>=95:
                        selected_fn=pfile
                        st="audio"
                        mycursor.execute("update coc_evidence set tamper_st=1 where id=%s",(maxid,))
                        mydb.commit()
                        break'''
                elif file_ext[1]=="mp4" or file_ext[1]=="avi":
                    d=1
                    hash1 = calculate_hash("static/upload1/"+pfile)
                    hash2 = calculate_hash("static/upload1/"+efile)

                    percentage_difference = hash_difference_percentage(hash1, hash2)
                    print(percentage_difference)
                    if percentage_difference>=95:
                        selected_fn=pfile
                        st="video"
                        x+=1
                        mycursor.execute("update coc_evidence set tamper_st=1,ftype='video',fname=%s,fid=%s where id=%s",(pfile,fid,maxid))
                        mydb.commit()
                        tamper_video(pfile,efile,"1")
                        break

        file1=selected_fn
        file2=efile

        if x>0:
            bdata1="ID:"+str(maxid)+", Case ID:"+case_id+", Status:Tamper Detected, uploaded by "+uname+", File: "+efile+", ID:"+str(maxid)+", Original: "+pfile+", ID:"+str(fid)
            cocchain(str(maxid),case_id,bdata1,'Tamper')

            mess="Case ID:"+case_id+", Tamper Detected, uploaded by "+uname+", File: "+efile+", ID:"+str(maxid)+", Original: "+pfile+", ID:"+str(fid)
            
            
        '''if st=="no":
            msg="verify"
        elif st=="docx":
            msg="docxverify"
        elif st=="pdf":
            msg="pdfverify"
        elif st=="audio":
            msg="audioverify"
        elif st=="video":
            msg="videoverify"
        else:'''
        msg="success"
        
        
        '''mycursor.execute('SELECT * FROM coc_evidence WHERE id=%s', (maxid,))
        dt = mycursor.fetchall()
        cutoff=10
        for rr in dt:
            hash0 = imagehash.average_hash(Image.open("static/upload1/"+rr[3])) 
            hash1 = imagehash.average_hash(Image.open("static/upload1/"+efile))
            cc1=hash0 - hash1
            print("cc="+str(cc1))
            if cc1<=cutoff:
                ss="ok"
                pre_id=str(rr[0])
                
                break
            else:
                ss="no"
        if ss=="ok":
            mycursor.execute('SELECT * FROM coc_evidence where id=%s',(maxid,))
            sp3 = mycursor.fetchone()
            dtime=str(sp3[4])
            mycursor.execute('SELECT * FROM coc_evidence where id=%s',(pre_id,))
            sp1 = mycursor.fetchone()
            pre_user=sp1[6]
            mycursor.execute('SELECT * FROM coc_register where uname=%s',(pre_user,))
            sp2 = mycursor.fetchone()
            pre_vid=sp2[0]
            

            bdata1="ID:"+str(pre_vid)+", Case ID:"+sp3[1]+", Status:Attack Found, Similar Evidence uploaded by "+uname+", Evidence ID:"+str(maxid)+", File: "+sp3[3]+" (Previous ID:"+str(pre_id)+")"
            cocchain(str(pre_vid),sp3[1],bdata1,'Attack')
            msg1="attack"

            mycursor.execute("SELECT max(id)+1 FROM coc_attack")
            maxid2 = mycursor.fetchone()[0]
            if maxid2 is None:
                maxid2=1
                
            sql2 = "INSERT INTO coc_attack(id,uname,eid,efile,case_id,status) VALUES (%s,%s,%s,%s,%s,%s)"
            val2 = (maxid2,uname,str(maxid),sp3[3],sp3[1],'0')
            mycursor.execute(sql2, val2)
            mydb.commit()'''
            
        ########
    ##########################
    if act=="req":
        mycursor.execute("SELECT * FROM coc_evidence where id=%s",(eid,))
        dd4 = mycursor.fetchone()
        mycursor.execute('SELECT count(*) FROM coc_access WHERE uname=%s && eid=%s', (uname,eid))
        c1 = mycursor.fetchone()[0]
        if c1==0:

            mycursor.execute("SELECT max(id)+1 FROM coc_access")
            maxid = mycursor.fetchone()[0]
            if maxid is None:
                maxid=1
                
            sql = "INSERT INTO coc_access(id,uname,eid,case_id,view_st,download_st) VALUES (%s,%s,%s,%s,%s,%s)"
            val = (maxid,uname,eid,case_id,'1','2')
            mycursor.execute(sql, val)
            mydb.commit()

            ###
            mycursor.execute('SELECT * FROM coc_access WHERE id=%s', (maxid,))
            dd = mycursor.fetchone()
            dtime=str(dd[6])
            bcdata="Evidence ID:"+eid+", Case ID:"+case_id+", Status: "+dd4[3]+", Download Request by "+uname
            cocchain(eid,case_id,bcdata,'Down-Request')
            ###
        else:
            mycursor.execute("update coc_access set view_st='1',download_st='2' where uname=%s && case_id=%s", (uname,case_id))
            mydb.commit()
            ###
            mycursor.execute("SELECT * FROM coc_access WHERE uname=%s && eid=%s", (uname,eid))
            dd = mycursor.fetchone()
            dtime=str(dd[6])
            bcdata="Evidence ID:"+eid+", Case ID:"+case_id+", Status: "+dd4[3]+", Download Request by "+uname
            cocchain(eid,case_id,bcdata,'Down-Request')
            ###

        mycursor.execute("SELECT max(id)+1 FROM coc_request")
        maxid2 = mycursor.fetchone()[0]
        if maxid2 is None:
            maxid2=1

        msgg="Evidence ID:"+eid+", Case ID:"+case_id+", File: "+dd4[3]+", Download Request by "+uname
        sql = "INSERT INTO coc_request(id,uname,message,reply,status) VALUES (%s,%s,%s,%s,%s)"
        val = (maxid2,uname,msgg,'','0')
        mycursor.execute(sql, val)
        mydb.commit()
        msg="req"



    #####################

        
    
    return render_template('a_upload.html',msg=msg,email=email,mess=mess,act=act,cid=cid,case_id=case_id,data1=data1,data2=data2,st=st,msg1=msg1,value=value,file1=file1,file2=file2)


@app.route('/a_attack', methods=['GET', 'POST'])
def a_attack():
    msg=""
    email=""
    mess=""
    efile=""
    case_id=""
    data1=[]
    data2=[]
    bdata1=""
    msg1=""
    st=""
    file1=""
    file2=""
    pfile=""
    fid=0
    fidd=""

    mycursor = mydb.cursor()

    mycursor.execute("SELECT * FROM coc_login where username='admin'")
    avalue = mycursor.fetchone()
    email=avalue[4]

    mycursor.execute('SELECT * FROM coc_evidence WHERE tamper_st=0 && attack_st=0')
    drow = mycursor.fetchall()

    x=0
    for drow1 in drow:
        case_id=drow1[1]
        pfile=drow1[3]
        fid=drow1[0]
        fidd=str(drow1[0])

        file_ext=pfile.split(".")
                   

        if file_ext[1]=="jpg" or file_ext[1]=="jpeg" or file_ext=="png":
            d=1

            hash1 = calculate_hash("static/assets/ud/"+pfile)
            hash2 = calculate_hash("static/upload1/"+pfile)

            percentage_difference = hash_difference_percentage(hash1, hash2)
            if percentage_difference>=95:
                selected_fn=pfile
                st="no"
                x+=1
                mycursor.execute("update coc_evidence set attack_st=1,ftype='image',hash2=%s where id=%s",(hash2,fid))
                mydb.commit()
                break
                
                
            
        elif file_ext[1]=="docx":
            d=1
            hash1 = calculate_hash("static/assets/ud/"+pfile)
            hash2 = calculate_hash("static/upload1/"+pfile)

            percentage_difference = hash_difference_percentage(hash1, hash2)
            if percentage_difference>=95:
                selected_fn=pfile
                st="docx"
                x+=1
                mycursor.execute("update coc_evidence set attack_st=1,ftype='docx',hash2=%s where id=%s",(hash2,fid))
                mydb.commit()
                break
        elif file_ext[1]=="pdf":
            d=1
            hash1 = calculate_hash("static/assets/ud/"+pfile)
            hash2 = calculate_hash("static/upload1/"+pfile)

            percentage_difference = hash_difference_percentage(hash1, hash2)
            if percentage_difference>=95:
                selected_fn=pfile
                st="pdf"
                x+=1
                mycursor.execute("update coc_evidence set attack_st=1,ftype='pdf',hash2=%s where id=%s",(hash2,fid))
                mydb.commit()
                break
        elif file_ext[1]=="wav" or file_ext[1]=="mp3":
            d=1
            #plot_waveform(pfile,str(fid))
            res=compare_audio("static/assets/ud/"+pfile,"static/upload1/"+pfile)
            print("audio per")
            print(res[2])
            if res[2]>0 and res[2]<3:
                
                st="audio"
                x+=1
                mycursor.execute("update coc_evidence set attack_st=1,ftype='audio',hash2=%s where id=%s",(hash2,fid))
                mydb.commit()
                break
            '''hash1 = calculate_hash("static/upload1/"+pfile)
            hash2 = calculate_hash("static/upload1/"+efile)

            percentage_difference = hash_difference_percentage(hash1, hash2)
            if percentage_difference>=95:
                selected_fn=pfile
                st="audio"
                mycursor.execute("update coc_evidence set tamper_st=1 where id=%s",(maxid,))
                mydb.commit()
                break'''
        elif file_ext[1]=="mp4" or file_ext[1]=="avi":
            d=1
            hash1 = calculate_hash("static/assets/ud/"+pfile)
            hash2 = calculate_hash("static/upload1/"+pfile)

            percentage_difference = hash_difference_percentage(hash1, hash2)
            print(percentage_difference)
            if percentage_difference>=95:
                selected_fn=pfile
                st="video"
                x+=1
                mycursor.execute("update coc_evidence set attack_st=1,ftype='video',hash2=%s where id=%s",(hash2,fid))
                mydb.commit()
                tamper_video(pfile,pfile,"2")
                break

    #file1=selected_fn
    #file2=efile

    if x>0:
        print("fid")
        print(case_id)
        print(pfile)
        
        mycursor.execute('SELECT * FROM coc_evidence WHERE id=%s',(fid,))
        drow2 = mycursor.fetchone()
        fidd=str(drow2[0])
        
        #bdata1="ID:"+fidd+", Case ID:"+case_id+", Status:File Attacked, File: "+pfile
        bdata1="ID: "+fidd+", Case ID: "+case_id+", Status: File Attacked, File: "+pfile
        cocchain(fidd,case_id,bdata1,'Attack')

        mess="Case ID:"+case_id+", File Attacked, File: "+pfile
        
    

    return render_template('a_attack.html',msg=msg,mess=mess,email=email,)


@app.route('/a_show', methods=['GET', 'POST'])
def a_show():
    msg=""
    st=""
    st2=""
    data=[]
    data3=[]

    file1=request.args.get("file1")
    file2=request.args.get("file2")
    uname=""
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM coc_evidence where filename=%s",(file1,))
    dat1 = mycursor.fetchone()

    mycursor.execute("SELECT * FROM coc_evidence where filename=%s",(file2,))
    dat2 = mycursor.fetchone()
    st="no"
    

    return render_template('a_show.html',msg=msg,dat1=dat1,dat2=dat2,st=st)

#CNN- Image Tamper Detection
def CNN():
    #Lets start by loading the Cifar10 data
    (X, y), (X_test, y_test) = cifar10.load_data()

    #Keep in mind the images are in RGB
    #So we can normalise the data by diving by 255
    #The data is in integers therefore we need to convert them to float first
    X, X_test = X.astype('float32')/255.0, X_test.astype('float32')/255.0

    #Then we convert the y values into one-hot vectors
    #The cifar10 has only 10 classes, thats is why we specify a one-hot
    #vector of width/class 10
    y, y_test = u.to_categorical(y, 10), u.to_categorical(y_test, 10)

    #Now we can go ahead and create our Convolution model
    model = Sequential()
    #We want to output 32 features maps. The kernel size is going to be
    #3x3 and we specify our input shape to be 32x32 with 3 channels
    #Padding=same means we want the same dimensional output as input
    #activation specifies the activation function
    model.add(Conv2D(32, (3, 3), input_shape=(32, 32, 3), padding='same',
                     activation='relu'))
    #20% of the nodes are set to 0
    model.add(Dropout(0.2))
    #now we add another convolution layer, again with a 3x3 kernel
    #This time our padding=valid this means that the output dimension can
    #take any form
    model.add(Conv2D(32, (3, 3), activation='relu', padding='valid'))
    #maxpool with a kernet of 2x2
    model.add(MaxPooling2D(pool_size=(2, 2)))
    #In a convolution NN, we neet to flatten our data before we can
    #input it into the ouput/dense layer
    model.add(Flatten())
    #Dense layer with 512 hidden units
    model.add(Dense(512, activation='relu'))
    #this time we set 30% of the nodes to 0 to minimize overfitting
    model.add(Dropout(0.3))
    #Finally the output dense layer with 10 hidden units corresponding to
    #our 10 classe
    model.add(Dense(10, activation='softmax'))
    #Few simple configurations
    model.compile(loss='categorical_crossentropy',
                  optimizer=SGD(momentum=0.5, decay=0.0004), metrics=['accuracy'])
    #Run the algorithm!
    model.fit(X, y, validation_data=(X_test, y_test), epochs=25,
              batch_size=512)
    #Save the weights to use for later
    model.save_weights("cifar10.hdf5")
    #Finally print the accuracy of our model!
    print("Accuracy: &2.f%%" %(model.evaluate(X_test, y_test)[1]*100))


@app.route('/a_tamper', methods=['GET', 'POST'])
def a_tamper():
    msg=""
    st=""
    st2=""
    data=[]
    data3=[]
    hash1=""
    hash2=""
    rtype=request.args.get("rtype")
    file1=request.args.get("file1")
    file2=request.args.get("file2")
    uname=""
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()

    
    mycursor.execute("SELECT * FROM coc_evidence where filename=%s",(file1,))
    dat1 = mycursor.fetchone()
    hash1=dat1[7]

    mycursor.execute("SELECT * FROM coc_evidence where filename=%s",(file2,))
    dat2 = mycursor.fetchone()
    if rtype=="1":
        hash2=dat2[13]
    else:
        hash2=dat2[7]
    
    st="no"

    ###################
    ff1=""
    if rtype=="1":
        ff1="static/assets/ud/"+file1
    else:
        ff1="static/assets/ud/"+file1
        
    before = cv2.imread(ff1)
    after = cv2.imread('static/upload1/'+file2)

    # Convert images to grayscale
    before_gray = cv2.cvtColor(before, cv2.COLOR_BGR2GRAY)
    after_gray = cv2.cvtColor(after, cv2.COLOR_BGR2GRAY)

    # Compute SSIM between the two images
    (score, diff) = structural_similarity(before_gray, after_gray, full=True)
    print("Image Similarity: {:.4f}%".format(score * 100))
    per=format(score * 100)

    # The diff image contains the actual image differences between the two images
    # and is represented as a floating point data type in the range [0,1] 
    # so we must convert the array to 8-bit unsigned integers in the range
    # [0,255] before we can use it with OpenCV
    diff = (diff * 255).astype("uint8")
    diff_box = cv2.merge([diff, diff, diff])

    # Threshold the difference image, followed by finding contours to
    # obtain the regions of the two input images that differ
    thresh = cv2.threshold(diff, 0, 255, cv2.THRESH_BINARY_INV | cv2.THRESH_OTSU)[1]
    contours = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    contours = contours[0] if len(contours) == 2 else contours[1]

    mask = np.zeros(before.shape, dtype='uint8')
    filled_after = after.copy()
    j=1
    for c in contours:
        area = cv2.contourArea(c)
        if area > 40:
            x,y,w,h = cv2.boundingRect(c)
            cv2.rectangle(before, (x, y), (x + w, y + h), (36,255,12), 2)
            mm=cv2.rectangle(after, (x, y), (x + w, y + h), (36,255,12), 2)
            cv2.imwrite("static/test/ggg.jpg", mm)

            image = cv2.imread("static/test/ggg.jpg")
            h1=h+10
            w1=w+30
            
            
            cropped = image[y:y+h1, x:x+w1]
            gg="static/test/f"+str(j)+".jpg"
            cv2.imwrite(""+gg, cropped)
        
            cv2.rectangle(diff_box, (x, y), (x + w, y + h), (36,255,12), 2)
            cv2.drawContours(mask, [c], 0, (255,255,255), -1)
            cv2.drawContours(filled_after, [c], 0, (0,255,0), -1)
            j+=1


    #print(j)
    ###################
    textarr=[]
    
    k=1
    while k<j:
        dt=[]
        ############
        fna="f"+str(k)+".jpg"
        print(fna)
        pytesseract.pytesseract.tesseract_cmd = 'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'
        Actual_image = cv2.imread("static/test/"+fna)
        #Sample_img = cv2.resize(Actual_image,(400,350))
        Image_ht,Image_wd,Image_thickness = Actual_image.shape
        Sample_img = cv2.cvtColor(Actual_image,cv2.COLOR_BGR2RGB)
        texts = pytesseract.image_to_data(Sample_img) 
        mytext=""
        prevy=0

        
        
        for cnt,text in enumerate(texts.splitlines()):
            
            if cnt==0:
                continue
            text = text.split()
            if len(text)==12:
                x,y,w,h = int(text[6]),int(text[7]),int(text[8]),int(text[9])
                if(len(mytext)==0):
                    prey=y
                if(prevy-y>=10 or y-prevy>=10):
                    #print(mytext)
                    s=1
                    #mytext=""
                mytext = mytext + text[11]+" "
                prevy=y

        v11=mytext
        print(v11)
        if v11=="":
            dt.append("")
            dt.append(fna)
        else:
            dt.append(v11)
            dt.append(fna)
            
        textarr.append(dt)
        
        k+=1

    #print(textarr)

    

    return render_template('a_tamper.html',msg=msg,dat1=dat1,dat2=dat2,st=st,textarr=textarr,hash1=hash1,hash2=hash2)


@app.route('/a_req', methods=['GET', 'POST'])
def a_req():
    msg=""
    st=""
    st2=""
    data=[]
    data3=[]
    uname=""
    if 'username' in session:
        uname = session['username']

    bdata=""
   
    
    mycursor = mydb.cursor()
    mycursor.execute("SELECT count(*) FROM coc_request where uname=%s order by id desc",(uname,))
    cnt = mycursor.fetchone()[0]
    if cnt>0:
        st="1"
        mycursor.execute("SELECT * FROM coc_request where uname=%s order by id desc",(uname,))
        data = mycursor.fetchall()

    mycursor.execute("SELECT count(*) FROM coc_request where cname=%s order by id desc",(uname,))
    cnt2 = mycursor.fetchone()[0]
    if cnt2>0:
        st2="1"
        mycursor.execute("SELECT * FROM coc_request where cname=%s order by id desc",(uname,))
        data3 = mycursor.fetchall()

    if request.method == 'POST':
        
        message = request.form['message']
        mycursor.execute("SELECT max(id)+1 FROM coc_request")
        maxid = mycursor.fetchone()[0]
        if maxid is None:
            maxid=1
        sql = "INSERT INTO coc_request(id,uname,message,reply,status) VALUES (%s,%s,%s,%s,%s)"
        val = (maxid,uname,message,'','0')
        mycursor.execute(sql, val)
        mydb.commit()
        ###
        mycursor.execute('SELECT * FROM coc_request WHERE id=%s', (maxid,))
        dd = mycursor.fetchone()
        dtime=str(dd[5])
        bcdata="Request ID:"+str(maxid)+", User ID:"+uname+", Status:Request"
        cocchain(str(maxid),case_id,bcdata,'Down-Request')
        ###
        msg="send"

    return render_template('a_req.html',msg=msg,data=data,data3=data3,st=st,st2=st2)

@app.route('/verify', methods=['GET', 'POST'])
def verify():
    msg=""
    cnt=0
    uname=""
    data1=[]
    mess=""
    act=request.args.get("act")

    if act=="1":
        ff=open("static/cocchain.json","r")
        fj=ff.read()
        ff.close()

        fjj=fj.split('}')

        nn=len(fjj)
        nn2=nn-2
        i=0
        fsn=""
        while i<nn-1:
            if i==nn2:
                fsn+=fjj[i]+"}"
            else:
                fsn+=fjj[i]+"},"
            i+=1
            
        #fjj1='},'.join(fjj)
        
        fj1="["+fsn+"]"
        

        #ff=open("static/crop.json","w")
        #ff.write(fj1)
        #ff.close()
        
        #dataframe = pd.read_json("static/coc.json", orient='values')
        
        
        #for ss in dataframe.values:
            
        #    if ss[4]=="Farmer" or ss[5]=="Payment" or ss[5]=="Claim":
                
        #        data1.append(ss)
        

    ################
    if act=="11":
        s1="1"
        ff=open("static/assets/js/d1.txt","r")
        ds=ff.read()
        ff.close()

        drow=ds.split("#|")
        
        i=0
        for dr in drow:
            
            dr1=dr.split("##")
            dt=[]
            #if "Register" in dr1[2]:
            
                
                
            dt.append(dr1[0])
            dt.append(dr1[1])
            dt.append(dr1[2])
            dt.append(dr1[3])
            #dt.append(dr1[4])
            if "Tamper" in dr1[2] or "Attack" in dr1[2]:
                dt.append("2")
            else:
                dt.append("1")
            data1.append(dt)
    else:
        ff=open("static/assets/js/d1.txt","r")
        ds=ff.read()
        ff.close()

        drow=ds.split("#|")
        
        i=0
        for dr in drow:
            
            dr1=dr.split("##")
            dt=[]
            #if "Register" in dr1[2]:
                
            dt.append(dr1[0])
            dt.append(dr1[1])
            dt.append(dr1[2])
            dt.append(dr1[3])
            #dt.append(dr1[4])
            data1.append(dt)
        
    
    return render_template('verify.html',msg=msg,act=act,data1=data1)


@app.route('/verify2', methods=['GET', 'POST'])
def verify2():
    msg=""
    cnt=0
    uname=""
    data1=[]
    mess=""
    act=request.args.get("act")

    if act=="1":
        ff=open("static/cocchain.json","r")
        fj=ff.read()
        ff.close()

        fjj=fj.split('}')

        nn=len(fjj)
        nn2=nn-2
        i=0
        fsn=""
        while i<nn-1:
            if i==nn2:
                fsn+=fjj[i]+"}"
            else:
                fsn+=fjj[i]+"},"
            i+=1
            
        #fjj1='},'.join(fjj)
        
        fj1="["+fsn+"]"
        

        #ff=open("static/crop.json","w")
        #ff.write(fj1)
        #ff.close()
        
        #dataframe = pd.read_json("static/coc.json", orient='values')
        
        
        #for ss in dataframe.values:
            
        #    if ss[4]=="Farmer" or ss[5]=="Payment" or ss[5]=="Claim":
                
        #        data1.append(ss)
        

    ################
    if act=="11":
        s1="1"
        ff=open("static/assets/js/d1.txt","r")
        ds=ff.read()
        ff.close()

        drow=ds.split("#|")
        
        i=0
        for dr in drow:
            
            dr1=dr.split("##")
            dt=[]
            #if "Register" in dr1[2]:
            
                
                
            dt.append(dr1[0])
            dt.append(dr1[1])
            dt.append(dr1[2])
            dt.append(dr1[3])
            #dt.append(dr1[4])
            if "Tamper" in dr1[2] or "Attack" in dr1[2]:
                dt.append("2")
            else:
                dt.append("1")
            data1.append(dt)
    else:
        ff=open("static/assets/js/d1.txt","r")
        ds=ff.read()
        ff.close()

        drow=ds.split("#|")
        
        i=0
        for dr in drow:
            
            dr1=dr.split("##")
            dt=[]
            #if "Register" in dr1[2]:
                
            dt.append(dr1[0])
            dt.append(dr1[1])
            dt.append(dr1[2])
            dt.append(dr1[3])
            #dt.append(dr1[4])
            data1.append(dt)
        
    
    return render_template('verify2.html',msg=msg,act=act,data1=data1)

@app.route('/down', methods=['GET', 'POST'])
def down():
    eid=request.args.get('eid')
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM coc_evidence where id=%s",(eid,))
    data = mycursor.fetchone()
    fn=data[3]
    ff="E"+eid+".hash"
    file = open('static/upload/'+ff, 'rb')
    byte = file.read()
    file.close()
      
    decodeit = open('static/down/'+fn, 'wb')
    decodeit.write(base64.b64decode((byte)))
    decodeit.close()
    
    path="static/down/"+fn


    
    return send_file(path, as_attachment=True)

##
#BERT-Word Document Embedding
def BERT():
    super(BERTLM, self).__init__()
    self.vocab = vocab
    self.embed_dim =embed_dim
    self.tok_embed = Embedding(self.vocab.size, embed_dim, self.vocab.padding_idx)
    self.pos_embed = LearnedPositionalEmbedding(embed_dim, device=local_rank)
    self.seg_embed = Embedding(2, embed_dim, None)

    self.out_proj_bias = nn.Parameter(torch.Tensor(self.vocab.size))

    self.layers = nn.ModuleList()
    for i in range(layers):
        self.layers.append(TransformerLayer(embed_dim, ff_embed_dim, num_heads, dropout))
    self.emb_layer_norm = LayerNorm(embed_dim)
    self.one_more = nn.Linear(embed_dim, embed_dim)
    self.one_more_layer_norm = LayerNorm(embed_dim)
    self.one_more_nxt_snt = nn.Linear(embed_dim, embed_dim) 
    self.nxt_snt_pred = nn.Linear(embed_dim, 1)
    self.dropout = dropout
    self.device = local_rank

    if approx == "none":
        self.approx = None
    elif approx == "adaptive":
        self.approx = nn.AdaptiveLogSoftmaxWithLoss(self.embed_dim, self.vocab.size, [10000, 20000, 200000])
    else:
        raise NotImplementedError("%s has not been implemented"%approx)
    self.reset_parameters()

def reset_parameters(self):
    nn.init.constant_(self.out_proj_bias, 0.)
    nn.init.constant_(self.nxt_snt_pred.bias, 0.)
    nn.init.constant_(self.one_more.bias, 0.)
    nn.init.constant_(self.one_more_nxt_snt.bias, 0.)
    nn.init.normal_(self.nxt_snt_pred.weight, std=0.02)
    nn.init.normal_(self.one_more.weight, std=0.02)
    nn.init.normal_(self.one_more_nxt_snt.weight, std=0.02)

def work(self, inp, seg=None, layers=None):
   
    if layers is not None:
        tot_layers = len(self.layers)
        for x in layers:
            if not (-tot_layers <= x < tot_layers):
                raise ValueError('layer %d out of range '%x)
        layers = [ (x+tot_layers if x <0 else x) for x in layers]
        max_layer_id = max(layers)
    
    seq_len, bsz = inp.size()
    if seg is None:
        seg = torch.zeros_like(inp)
    x = self.tok_embed(inp) + self.seg_embed(seg) + self.pos_embed(inp)
    x = self.emb_layer_norm(x)
    x = F.dropout(x, p=self.dropout, training=self.training)
    padding_mask = torch.eq(inp, self.vocab.padding_idx)
    if not padding_mask.any():
        padding_mask = None
    
    xs = []
    for layer_id, layer in enumerate(self.layers):
        x, _ ,_ = layer(x, self_padding_mask=padding_mask)
        xs.append(x)
        if layers is not None and layer_id >= max_layer_id:
            break
    
    if layers is not None:
        x = torch.stack([xs[i] for i in layers])
        z = torch.tanh(self.one_more_nxt_snt(x[:,0,:,:]))
    else:
        z = torch.tanh(self.one_more_nxt_snt(x[0]))
    return x, z

def forward(self, truth, inp, seg, msk, nxt_snt_flag):
    seq_len, bsz = inp.size()
    x = self.tok_embed(inp) + self.seg_embed(seg) + self.pos_embed(inp)
    x = self.emb_layer_norm(x)
    x = F.dropout(x, p=self.dropout, training=self.training)
    padding_mask = torch.eq(truth, self.vocab.padding_idx)
    if not padding_mask.any():
        padding_mask = None
    for layer in self.layers:
        x, _ ,_ = layer(x, self_padding_mask=padding_mask)

    masked_x = x.masked_select(msk.unsqueeze(-1))
    masked_x = masked_x.view(-1, self.embed_dim)
    gold = truth.masked_select(msk)
    
    y = self.one_more_layer_norm(gelu(self.one_more(masked_x)))
    out_proj_weight = self.tok_embed.weight

    if self.approx is None:
        log_probs = torch.log_softmax(F.linear(y, out_proj_weight, self.out_proj_bias), -1)
    else:
        log_probs = self.approx.log_prob(y)

    loss = F.nll_loss(log_probs, gold, reduction='mean')

    z = torch.tanh(self.one_more_nxt_snt(x[0]))
    nxt_snt_pred = torch.sigmoid(self.nxt_snt_pred(z).squeeze(1))
    nxt_snt_acc = torch.eq(torch.gt(nxt_snt_pred, 0.5), nxt_snt_flag).float().sum().item()
    nxt_snt_loss = F.binary_cross_entropy(nxt_snt_pred, nxt_snt_flag.float(), reduction='mean')
    
    tot_loss = loss + nxt_snt_loss
    
    _, pred = log_probs.max(-1)
    tot_tokens = msk.float().sum().item()
    acc = torch.eq(pred, gold).float().sum().item()
    
    return (pred, gold), tot_loss, acc, tot_tokens, nxt_snt_acc, bsz
##################
from docx import Document
from diff_match_patch import diff_match_patch

#pdf
#import fitz

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"
    return full_text

def highlight_differences(text1, text2):
    dmp = diff_match_patch()
    diffs = dmp.diff_main(text1, text2)
    dmp.diff_cleanupSemantic(diffs)
    diff_html = dmp.diff_prettyHtml(diffs)
    return diff_html

@app.route('/result', methods=['GET', 'POST'])
def compare_files():
    text1=""
    hash1=""
    hash2=""
    uname=""
    if 'username' in session:
        uname = session['username']

    bdata=""
    rtype=request.args.get("rtype")
    f1=request.args.get("file1")
    f2=request.args.get("file2")
    
    mycursor = mydb.cursor()

    mycursor.execute("SELECT * FROM coc_evidence where filename=%s",(f1,))
    fdata1 = mycursor.fetchone()
    hash1=fdata1[7]

    mycursor.execute("SELECT * FROM coc_evidence where filename=%s",(f2,))
    fdata2 = mycursor.fetchone()
    if rtype=="1":
        hash2=fdata2[13]
    else:
        hash2=fdata2[7]
    
    '''if request.method == 'POST':
        file1 = request.files['file1']
        file2 = request.files['file2']'''
    file1=""
    if rtype=="1":
        file1="static/assets/ud/"+f1
    else:
        file1="static/assets/ud/"+f1
    file2="static/upload1/"+f2
    
    text1 = extract_text_from_docx(file1)
    text2 = extract_text_from_docx(file2)

    differences = highlight_differences(text1, text2)
    print("docx diff")
    print(differences)

    return render_template('result.html', differences=differences,text1=text1,hash1=hash1,hash2=hash2)
##PDF####################
def remove_img_on_pdf(idoc, page):
    img_list = idoc.getPageImageList(page)
    con_list = idoc[page].get_contents()

    for i in con_list:
        c = idoc.xref_stream(i)
        if c != None:
            for v in img_list:
                arr = bytes(v[7], 'utf-8')
                r = c.find(arr)
                if r != -1:
                    cnew = c.replace(arr, b"")
                    idoc.update_stream(i, cnew)
                    c = idoc.xref_stream(i)
    return idoc
def extract_text_from_pdf(file_path):

    #doc = fitz.open(file_path)
    #rdoc = remove_img_on_pdf(doc, 0)
    #rdoc.save(file_path)

    doc = fitz.open(file_path)
    full_text = ""
    for page in doc:
        full_text += page.get_text()
    return full_text

@app.route('/result2', methods=['GET', 'POST'])
def compare_files2():
    '''if request.method == 'POST':
        file1 = request.files['file1']
        file2 = request.files['file2']'''
    text1=""
    hash1=""
    hash2=""
    uname=""
    if 'username' in session:
        uname = session['username']

    bdata=""
    rtype=request.args.get("rtype")
    f1=request.args.get("file1")
    f2=request.args.get("file2")
    
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM coc_evidence where filename=%s",(f1,))
    fdata1 = mycursor.fetchone()
    hash1=fdata1[7]

    mycursor.execute("SELECT * FROM coc_evidence where filename=%s",(f2,))
    fdata2 = mycursor.fetchone()
    if rtype=="1":
        hash2=fdata2[13]
    else:
        hash2=fdata2[7]

    file1=""
    if rtype=="1":
        file1="static/assets/ud/"+f1
    else:
        file1="static/assets/ud/"+f1

    print(file1)
    file2="static/upload1/"+f2
    text1 = extract_text_from_pdf(file1)
    text2 = extract_text_from_pdf(file2)
    
    differences = highlight_differences(text1, text2)
    print("pfd diff")
    print(differences)

    return render_template('result2.html', differences=differences,text1=text1,hash1=hash1,hash2=hash2)

###Video###############
#TCN  - Temporal Convolutional Network - Video Frame Level Analysis
def is_power_of_two(num: int):
    return num != 0 and ((num & (num - 1)) == 0)


def adjust_dilations(dilations: list):
    if all([is_power_of_two(i) for i in dilations]):
        return dilations
    else:
        new_dilations = [2 ** i for i in dilations]
        return new_dilations


    

    def __init__(self,
                 dilation_rate: int,
                 nb_filters: int,
                 kernel_size: int,
                 padding: str,
                 activation: str = 'relu',
                 dropout_rate: float = 0,
                 kernel_initializer: str = 'he_normal',
                 use_batch_norm: bool = False,
                 use_layer_norm: bool = False,
                 use_weight_norm: bool = False,
                 **kwargs):

        self.dilation_rate = dilation_rate
        self.nb_filters = nb_filters
        self.kernel_size = kernel_size
        self.padding = padding
        self.activation = activation
        self.dropout_rate = dropout_rate
        self.use_batch_norm = use_batch_norm
        self.use_layer_norm = use_layer_norm
        self.use_weight_norm = use_weight_norm
        self.kernel_initializer = kernel_initializer
        self.layers = []
        self.shape_match_conv = None
        self.res_output_shape = None
        self.final_activation = None

        super(ResidualBlock, self).__init__(**kwargs)

    def tcn_full_summary(model: Model, expand_residual_blocks=True):
        #import tensorflow as tf
        # 2.6.0-rc1, 2.5.0...
        versions = [int(v) for v in tf.__version__.split('-')[0].split('.')]
        if versions[0] <= 2 and versions[1] < 5:
            layers = model._layers.copy()  # store existing layers
            model._layers.clear()  # clear layers

            for i in range(len(layers)):
                if isinstance(layers[i], TCN):
                    for layer in layers[i]._layers:
                        if not isinstance(layer, ResidualBlock):
                            if not hasattr(layer, '__iter__'):
                                model._layers.append(layer)
                        else:
                            if expand_residual_blocks:
                                for lyr in layer._layers:
                                    if not hasattr(lyr, '__iter__'):
                                        model._layers.append(lyr)
                            else:
                                model._layers.append(layer)
                else:
                    model._layers.append(layers[i])

            model.summary()  # print summary

            # restore original layers
            model._layers.clear()
            [model._layers.append(lyr) for lyr in layers]

            

        def _build_layer(self, layer):
           
            self.layers.append(layer)
            self.layers[-1].build(self.res_output_shape)
            self.res_output_shape = self.layers[-1].compute_output_shape(self.res_output_shape)

        def build(self, input_shape):

            with K.name_scope(self.name):  # name scope used to make sure weights get unique names
                self.layers = []
                self.res_output_shape = input_shape

                for k in range(2):  # dilated conv block.
                    name = 'conv1D_{}'.format(k)
                    with K.name_scope(name):  # name scope used to make sure weights get unique names
                        conv = Conv1D(
                            filters=self.nb_filters,
                            kernel_size=self.kernel_size,
                            dilation_rate=self.dilation_rate,
                            padding=self.padding,
                            name=name,
                            kernel_initializer=self.kernel_initializer
                        )
                        if self.use_weight_norm:
                            from tensorflow_addons.layers import WeightNormalization
                            # wrap it. WeightNormalization API is different than BatchNormalization or LayerNormalization.
                            with K.name_scope('norm_{}'.format(k)):
                                conv = WeightNormalization(conv)
                        self._build_layer(conv)

                    with K.name_scope('norm_{}'.format(k)):
                        if self.use_batch_norm:
                            self._build_layer(BatchNormalization())
                        elif self.use_layer_norm:
                            self._build_layer(LayerNormalization())
                        elif self.use_weight_norm:
                            pass  # done above.

                    with K.name_scope('act_and_dropout_{}'.format(k)):
                        self._build_layer(Activation(self.activation, name='Act_Conv1D_{}'.format(k)))
                        self._build_layer(SpatialDropout1D(rate=self.dropout_rate, name='SDropout_{}'.format(k)))

                if self.nb_filters != input_shape[-1]:
                    # 1x1 conv to match the shapes (channel dimension).
                    name = 'matching_conv1D'
                    with K.name_scope(name):
                        # make and build this layer separately because it directly uses input_shape.
                        # 1x1 conv.
                        self.shape_match_conv = Conv1D(
                            filters=self.nb_filters,
                            kernel_size=1,
                            padding='same',
                            name=name,
                            kernel_initializer=self.kernel_initializer
                        )
                else:
                    name = 'matching_identity'
                    self.shape_match_conv = Lambda(lambda x: x, name=name)

                with K.name_scope(name):
                    self.shape_match_conv.build(input_shape)
                    self.res_output_shape = self.shape_match_conv.compute_output_shape(input_shape)

                self._build_layer(Activation(self.activation, name='Act_Conv_Blocks'))
                self.final_activation = Activation(self.activation, name='Act_Res_Block')
                self.final_activation.build(self.res_output_shape)  # probably isn't necessary

                # this is done to force Keras to add the layers in the list to self._layers
                for layer in self.layers:
                    self.__setattr__(layer.name, layer)
                self.__setattr__(self.shape_match_conv.name, self.shape_match_conv)
                self.__setattr__(self.final_activation.name, self.final_activation)

                super(ResidualBlock, self).build(input_shape)  # done to make sure self.built is set True

        def call(self, inputs, training=None, **kwargs):
            """
            Returns: A tuple where the first element is the residual model tensor, and the second
                     is the skip connection tensor.
            """
            
            x1 = inputs
            for layer in self.layers:
                training_flag = 'training' in dict(inspect.signature(layer.call).parameters)
                x1 = layer(x1, training=training) if training_flag else layer(x1)
            x2 = self.shape_match_conv(inputs)
            x1_x2 = self.final_activation(layers.add([x2, x1], name='Add_Res'))
            return [x1_x2, x1]

        def compute_output_shape(self, input_shape):
            return [self.res_output_shape, self.res_output_shape]


'''def detect_tampering(video_path1, video_path2, output_dir):
    cap1 = cv2.VideoCapture(video_path1)
    cap2 = cv2.VideoCapture(video_path2)

    frame_number = 0
    tampering_frames = []

    while True:
        ret1, frame1 = cap1.read()
        ret2, frame2 = cap2.read()

        if not ret1 or not ret2:
            break

        # Convert frames to grayscale for frame differencing
        gray_frame1 = cv2.cvtColor(frame1, cv2.COLOR_BGR2GRAY)
        gray_frame2 = cv2.cvtColor(frame2, cv2.COLOR_BGR2GRAY)

        # Compute absolute difference between frames
        diff = cv2.absdiff(gray_frame1, gray_frame2)

        # Check if tampering detected based on a threshold
        threshold = 30
        if cv2.countNonZero(diff) > threshold:
            tampering_frames.append(frame_number)

            # Save tampering frame as an image
            output_path = os.path.join(output_dir, f"tampering_frame_{frame_number}.jpg")
            cv2.imwrite(output_path, frame1)

        frame_number += 1

    return tampering_frames'''

@app.route('/result_video', methods=['GET', 'POST'])
def compare_videos():
    '''if request.method == 'POST':
        video1 = request.files['video1']
        video2 = request.files['video2']'''
    uname=""
    hash1=""
    hash2=""
    if 'username' in session:
        uname = session['username']

    bdata=""
    '''file1=request.args.get("file1")
    file2=request.args.get("file2")
    
    mycursor = mydb.cursor()
    
    video1="static/upload1/"+file1
    video2="static/upload1/"+file2

    # Create a temporary directory to store tampering frames
    output_dir = "static/temp_output"
    os.makedirs(output_dir, exist_ok=True)

    tampering_frames = detect_tampering(video1, video2, output_dir)'''
    act=request.args.get("act")
    rtype=request.args.get("rtype")
    f1=request.args.get("file1")
    f2=request.args.get("file2")

    fv1="static/assets/ud/"+f1
    fv2="static/upload1/"+f2
    
    ff=open("file1.txt","w")
    ff.write(fv1)
    ff.close()
    ff=open("file2.txt","w")
    ff.write(fv2)
    ff.close()

    tamper_video(f1,f2,"1")
    
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM coc_evidence where filename=%s",(f1,))
    fdata1 = mycursor.fetchone()
    hash1=fdata1[7]

    mycursor.execute("SELECT * FROM coc_evidence where filename=%s",(f2,))
    fdata2 = mycursor.fetchone()
    if rtype=="1":
        hash2=fdata2[13]
    else:
        hash2=fdata2[7]

    dimg1=[]
    dimg2=[]
    folder = 'static/video3'
    for filename in os.listdir(folder):
        dt1=[]
        dt2=[]
        file_path = os.path.join(folder, filename)

        af=filename.split(".")
        af1=af[0]
        af2=len(af1)
        fr_no=af1[1:af2]

        dt1.append(fr_no)
        dt1.append(filename)

        dimg2.append(dt1)

        n=len(filename)
        fn2=filename[1:n]
        fn3="a"+fn2

        dt2.append(fr_no)
        dt2.append(fn3)
        dimg1.append(dt2)
    

    return render_template('result_video.html',dimg1=dimg1,dimg2=dimg2,hash1=hash1,hash2=hash2,file1=f1,file2=f2,rtype=rtype,act=act)
####
def plot_waveform(audio_file,fid):

    fn="g"+fid+".png"
    with wave.open("static/upload1/"+audio_file, 'r') as wav_file:
        # Extracting raw audio data
        frames = wav_file.readframes(-1)
        signal = np.frombuffer(frames, dtype='int16')

        # Get the frame rate (sampling frequency) and duration of the audio
        framerate = wav_file.getframerate()
        duration = len(signal) / framerate

        # Time axis for the waveform
        time = np.linspace(0, duration, num=len(signal))

        # Plotting
        plt.figure(figsize=(10, 6))
        plt.plot(time, signal, color='blue')
        #plt.title('Waveform of {}'.format(audio_file))
        plt.title("Waveform")
        plt.xlabel('Time (s)')
        plt.ylabel('Amplitude')
        plt.grid(True)
        plt.savefig('static/graph/'+fn)
        #plt.show()
##############################
#Hidden Markov Model-Audio Spectrum Analysis
def HMM():
    self.transmission_prob = transmission_prob
    self.emission_prob = emission_prob
    self.n = self.emission_prob.shape[1]
    self.m = self.emission_prob.shape[0]
    self.observations = None
    self.forward = []
    self.backward = []
    self.psi = []
    self.obs = obs
    self.emiss_ref = {}
    self.forward_final = [0 , 0]
    self.backward_final = [0 , 0]
    self.state_probs = []
    if obs is None and self.observations is not None:
        self.obs = self.assume_obs()

    def assume_obs(self):
        '''
        If observation labels are not given, will assume that the emission
        probabilities are in alpha-numerical order.
        '''
        obs = list(set(list(self.observations)))
        obs.sort()
        for i in range(len(obs)):
            self.emiss_ref[obs[i]] = i
        return obs

    def train(self, observations, iterations = 10, verbose=True):
        '''
        Trains the model parameters according to the observation sequence.

        Input:
        - observations: 1-D string array of T observations
        '''
        self.observations = observations
        self.obs = self.assume_obs()
        self.psi = [[[0.0] * (len(self.observations)-1) for i in range(self.n)] for i in range(self.n)]
        self.gamma = [[0.0] * (len(self.observations)) for i in range(self.n)]
        for i in range(iterations):
            old_transmission = self.transmission_prob.copy()
            old_emission = self.emission_prob.copy()
            if verbose:
                print("Iteration: {}".format(i + 1))
            self.expectation()
            self.maximization()

    def expectation(self):
        '''
        Executes expectation step.
        '''
        self.forward = self.forward_recurse(len(self.observations))
        self.backward = self.backward_recurse(0)
        self.get_gamma()
        self.get_psi()

    def get_gamma(self):
        '''
        Calculates the gamma matrix.
        '''
        self.gamma = [[0, 0] for i in range(len(self.observations))]
        for i in range(len(self.observations)):
            self.gamma[i][0] = (float(self.forward[0][i] * self.backward[0][i]) /
                                float(self.forward[0][i] * self.backward[0][i] +
                                self.forward[1][i] * self.backward[1][i]))
            self.gamma[i][1] = (float(self.forward[1][i] * self.backward[1][i]) /
                                float(self.forward[0][i] * self.backward[0][i] +
                                self.forward[1][i] * self.backward[1][i]))

    def get_psi(self):
        '''
        Runs the psi calculation.
        '''
        for t in range(1, len(self.observations)):
            for j in range(self.n):
                for i in range(self.n):
                    self.psi[i][j][t-1] = self.calculate_psi(t, i, j)

    def calculate_psi(self, t, i, j):
        '''
        Calculates the psi for a transition from i->j for t > 0.
        '''
        alpha_tminus1_i = self.forward[i][t-1]
        a_i_j = self.transmission_prob[j+1][i+1]
        beta_t_j = self.backward[j][t]
        observation = self.observations[t]
        b_j = self.emission_prob[self.emiss_ref[observation]][j]
        denom = float(self.forward[0][i] * self.backward[0][i] + self.forward[1][i] * self.backward[1][i])
        return (alpha_tminus1_i * a_i_j * beta_t_j * b_j) / denom

    def maximization(self):
        '''
        Executes maximization step.
        '''
        self.get_state_probs()
        for i in range(self.n):
            self.transmission_prob[i+1][0] = self.gamma[0][i]
            self.transmission_prob[-1][i+1] = self.gamma[-1][i] / self.state_probs[i]
            for j in range(self.n):
                self.transmission_prob[j+1][i+1] = self.estimate_transmission(i, j)
            for obs in range(self.m):
                self.emission_prob[obs][i] = self.estimate_emission(i, obs)

    def get_state_probs(self):
        '''
        Calculates total probability of a given state.
        '''
        self.state_probs = [0] * self.n
        for state in range(self.n):
            summ = 0
            for row in self.gamma:
                summ += row[state]
            self.state_probs[state] = summ

    def estimate_transmission(self, i, j):
        '''
        Estimates transmission probabilities from i to j.
        '''
        return sum(self.psi[i][j]) / self.state_probs[i]

    def estimate_emission(self, j, observation):
        '''
        Estimate emission probability for an observation from state j.
        '''
        observation = self.obs[observation]
        ts = [i for i in range(len(self.observations)) if self.observations[i] == observation]
        for i in range(len(ts)):
            ts[i] = self.gamma[ts[i]][j]
        return sum(ts) / self.state_probs[j]

    def backward_recurse(self, index):
        '''
        Runs the backward recursion.
        '''
        # Initialization at T
        if index == (len(self.observations) - 1):
            backward = [[0.0] * (len(self.observations)) for i in range(self.n)]
            for state in range(self.n):
                backward[state][index] = self.backward_initial(state)
            return backward
        # Recursion for T --> 0
        else:
            backward = self.backward_recurse(index+1)
            for state in range(self.n):
                if index >= 0:
                    backward[state][index] = self.backward_probability(index, backward, state)
                if index == 0:
                    self.backward_final[state] = self.backward_probability(index, backward, 0, final=True)
            return backward

    def backward_initial(self, state):
        '''
        Initialization of backward probabilities.
        '''
        return self.transmission_prob[self.n + 1][state + 1]

    def backward_probability(self, index, backward, state, final=False):
        '''
        Calculates the backward probability at index = t.
        '''
        p = [0] * self.n
        for j in range(self.n):
            observation = self.observations[index + 1]
            if not final:
                a = self.transmission_prob[j + 1][state + 1]
            else:
                a = self.transmission_prob[j + 1][0]
            b = self.emission_prob[self.emiss_ref[observation]][j]
            beta = backward[j][index + 1]
            p[j] = a * b * beta
        return sum(p)

    def forward_recurse(self, index):
        '''
        Executes forward recursion.
        '''
        # Initialization
        if index == 0:
            forward = [[0.0] * (len(self.observations)) for i in range(self.n)]
            for state in range(self.n):
                forward[state][index] = self.forward_initial(self.observations[index], state)
            return forward
        # Recursion
        else:
            forward = self.forward_recurse(index-1)
            for state in range(self.n):
                if index != len(self.observations):
                    forward[state][index] = self.forward_probability(index, forward, state)
                else:
                    # Termination
                    self.forward_final[state] = self.forward_probability(index, forward, state, final=True)
            return forward

    def forward_initial(self, observation, state):
        '''
        Calculates initial forward probabilities.
        '''
        self.transmission_prob[state + 1][0]
        self.emission_prob[self.emiss_ref[observation]][state]
        return self.transmission_prob[state + 1][0] * self.emission_prob[self.emiss_ref[observation]][state]

    def forward_probability(self, index, forward, state, final=False):
        '''
        Calculates the alpha for index = t.
        '''
        p = [0] * self.n
        for prev_state in range(self.n):
            if not final:
                # Recursion
                obs_index = self.emiss_ref[self.observations[index]]
                p[prev_state] = forward[prev_state][index-1] * self.transmission_prob[state + 1][prev_state + 1] * self.emission_prob[obs_index][state]
            else:
                # Termination
                p[prev_state] = forward[prev_state][index-1] * self.transmission_prob[self.n][prev_state + 1]
        return sum(p)

    def likelihood(self, new_observations):
        '''
        Returns the probability of a observation sequence based on current model
        parameters.
        '''
        new_hmm = HMM(self.transmission_prob, self.emission_prob)
        new_hmm.observations = new_observations
        new_hmm.obs = new_hmm.assume_obs()
        forward = new_hmm.forward_recurse(len(new_observations))
        return sum(new_hmm.forward_final)


    model = HMM(transmission, emission)
    model.train(observations)
    print("Model transmission probabilities:\n{}".format(model.transmission_prob))
    print("Model emission probabilities:\n{}".format(model.emission_prob))

###
#import librosa
#from pydub import AudioSegment

def compare_audio(audio_path1, audio_path2):
    # Load audio files
    audio1, sr1 = librosa.load(audio_path1, sr=None)
    audio2, sr2 = librosa.load(audio_path2, sr=None)

    # Resample if sample rates are different
    if sr1 != sr2:
        audio2 = librosa.resample(audio2, sr2, sr1)
        sr2 = sr1

    # Compute waveforms
    waveform1, _ = librosa.effects.trim(audio1)
    waveform2, _ = librosa.effects.trim(audio2)

    # Ensure both waveforms have the same length
    min_length = min(len(waveform1), len(waveform2))
    waveform1 = waveform1[:min_length]
    waveform2 = waveform2[:min_length]

    # Calculate difference between waveforms
    difference = np.abs(waveform1 - waveform2)

    # Calculate percentage difference
    percentage_difference = np.mean(difference) / np.max(np.abs(waveform1)) * 100

    return waveform1, waveform2, percentage_difference

@app.route('/result_audio', methods=['GET', 'POST'])
def compare_audio_files():
    '''if request.method == 'POST':
        audio1 = request.files['audio1']
        audio2 = request.files['audio2']

        audio1_path = 'temp_audio/audio1.wav'
        audio2_path = 'temp_audio/audio2.wav'

        audio1.save(audio1_path)
        audio2.save(audio2_path)'''
    hash1=""
    hash2=""
    gfile=""
    uname=""
    if 'username' in session:
        uname = session['username']

    bdata=""
    act=request.args.get("act")
    rtype=request.args.get("rtype")
    file1=request.args.get("file1")
    file2=request.args.get("file2")
    
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM coc_evidence where filename=%s",(file1,))
    fdata1 = mycursor.fetchone()
    hash1=fdata1[7]
    gfile="g"+str(fdata1[0])+".png"

    mycursor.execute("SELECT * FROM coc_evidence where filename=%s",(file2,))
    fdata2 = mycursor.fetchone()
    if rtype=="1":
        hash2=fdata2[13]
    else:
        hash2=fdata2[7]
    

    

    audio1_path=""
    if rtype=="1":
        audio1_path="static/assets/ud/"+file1
    else:
        audio1_path="static/assets/ud/"+file1
        
    audio2_path="static/upload1/"+file2
    waveform1, waveform2, percentage_difference = compare_audio(audio1_path, audio2_path)

    # Plot waveforms
    plt.figure(figsize=(10, 6))
    plt.plot(waveform1, label='Audio 1')
    plt.plot(waveform2, label='Audio 2')
    plt.title('Audio Waveform Comparison')
    plt.xlabel('Time')
    plt.ylabel('Amplitude')
    plt.legend()
    plt.savefig('static/waveform_comparison.png')

    return render_template('result_audio.html', percentage_difference=percentage_difference,hash1=hash1,hash2=hash2,gfile=gfile,file1=file1,file2=file2,rtype=rtype,act=act)


@app.route('/ta_home', methods=['GET', 'POST'])
def ta_home():
    msg=""
    act=request.args.get("act")
    email=""
    mess=""
    bdata=""
    
    
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM coc_case order by id desc")
    data = mycursor.fetchall()

   
    
    return render_template('ta_home.html',msg=msg,act=act,data=data)

@app.route('/ta_tamper', methods=['GET', 'POST'])
def ta_tamper():
    msg=""
    act=request.args.get("act")
    cid=request.args.get("cid")
    email=""
    mess=""
    st=""
    bdata=""
    x=0
    
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM coc_case where id=%s",(cid,))
    dd = mycursor.fetchone()
    case_id=dd[1]
    
    mycursor.execute("SELECT * FROM coc_evidence where case_id=%s order by id desc",(case_id,))
    data = mycursor.fetchall()

    mycursor.execute("SELECT * FROM coc_evidence where attack_st=1 order by id desc")
    data2 = mycursor.fetchall()

    for d2 in data2:
        x+=1
    if x>0:
        st="1"
    
    return render_template('ta_tamper.html',msg=msg,act=act,data=data,data2=data2,st=st)




@app.route('/logout')
def logout():
    # remove the username from the session if it is there
    session.pop('username', None)
    return redirect(url_for('index'))

def gen2(camera):
    
    while True:
        frame = camera.get_frame()
        
        
        yield (b'--frame\r\n'
               b'Content-Type: image/jpeg\r\n\r\n' + frame + b'\r\n\r\n')
    
@app.route('/video_feed2')
        

def video_feed2():
    return Response(gen2(VideoCamera2()),
                    mimetype='multipart/x-mixed-replace; boundary=frame')
#######
def gen(camera):
    
    while True:
        frame = camera.get_frame()
        
        
        yield (b'--frame\r\n'
               b'Content-Type: image/jpeg\r\n\r\n' + frame + b'\r\n\r\n')
    
@app.route('/video_feed')
        

def video_feed():
    return Response(gen(VideoCamera()),
                    mimetype='multipart/x-mixed-replace; boundary=frame')

if __name__ == "__main__":
    app.secret_key = os.urandom(12)
    app.run(debug=True,host='0.0.0.0', port=5000)
